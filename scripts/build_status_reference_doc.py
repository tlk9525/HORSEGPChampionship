from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "horse-racing-status-reference.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)

    for margin, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "D9E2EC")


def set_table_width(table, width_dxa=9360, indent_dxa=120):
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def set_column_widths(table, widths):
    table.autofit = False

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def style_paragraph(paragraph, size=11, color="000000", bold=False):
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor.from_string(color)
        run.bold = bold
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.25


def add_table(doc, rows, widths=(1900, 7460)):
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(table)
    set_table_borders(table)
    set_column_widths(table, widths)

    header = table.rows[0].cells
    for index, value in enumerate(rows[0]):
        header[index].text = value
        set_cell_shading(header[index], "E8EEF5")
        set_cell_margins(header[index])
        header[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in header[index].paragraphs:
            style_paragraph(paragraph, size=10, color="0B2545", bold=True)

    for row_values in rows[1:]:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            cells[index].text = value
            set_cell_margins(cells[index])
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cells[index].paragraphs:
                style_paragraph(paragraph, size=10, color="111827", bold=index == 0)

    doc.add_paragraph()
    return table


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_title(doc):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run("Bảng trạng thái hệ thống Horse Racing")
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string("0B2545")

    subtitle = doc.add_paragraph(
        "Tài liệu tham khảo các trạng thái đang dùng trong schema, seed và backend."
    )
    style_paragraph(subtitle, size=11, color="4B5563")


def main():
    doc = Document()
    configure_document(doc)
    add_title(doc)

    doc.add_heading("Tóm tắt luồng chính", level=1)
    add_table(
        doc,
        [
            ("Thành phần", "Luồng trạng thái chính"),
            ("User", "pending -> active hoặc rejected"),
            ("Tournament", "registration -> approvals -> active -> completed"),
            ("Horse", "pending -> approved hoặc rejected; có thể retired"),
            ("Horse + Jockey", "pending-jockey -> pending-admin -> confirmed"),
            ("Race", "registration-open -> registration-closed -> published -> in-progress -> finished -> completed"),
            ("Race Entry", "pending-approval -> approved hoặc rejected; kết quả: draft -> official"),
        ],
        widths=(2300, 7060),
    )

    sections = [
        (
            "users",
            "Trạng thái tài khoản người dùng.",
            [
                ("pending", "Chờ Admin duyệt tài khoản."),
                ("active", "Tài khoản hoạt động, được đăng nhập."),
                ("approved", "Có trong schema, dùng như trạng thái duyệt tổng quát."),
                ("rejected", "Tài khoản bị từ chối."),
                ("suspended", "Tạm khóa, hiện ít dùng trong luồng chính."),
                ("locked", "Khóa tài khoản, hiện ít dùng trong luồng chính."),
            ],
        ),
        (
            "tournaments",
            "Trạng thái của giải đấu.",
            [
                ("registration", "Giải đã tạo và đang mở đăng ký."),
                ("approvals", "Giai đoạn Admin duyệt hồ sơ Owner/Jockey."),
                ("active", "Giải đang hoạt động."),
                ("completed", "Giải đã hoàn tất."),
            ],
        ),
        (
            "horses",
            "Trạng thái hồ sơ ngựa do Owner đăng ký.",
            [
                ("draft", "Bản nháp, có trong schema nhưng ít dùng."),
                ("pending", "Owner gửi hồ sơ ngựa, chờ Admin duyệt."),
                ("approved", "Ngựa được Admin duyệt."),
                ("rejected", "Ngựa bị Admin từ chối."),
                ("retired", "Ngựa ngừng thi đấu."),
            ],
        ),
        (
            "horses.jockeyConfirmation",
            "Trạng thái ghép jockey cho ngựa.",
            [
                ("waiting-owner", "Chưa có lựa chọn hoặc chờ Owner chọn jockey."),
                ("pending", "Đã gửi lời mời/chờ xử lý."),
                ("pending-jockey", "Chờ Jockey chấp nhận lời mời."),
                ("pending-admin", "Jockey đã chấp nhận, chờ Admin duyệt cặp."),
                ("confirmed", "Admin đã duyệt cặp Horse + Jockey."),
            ],
        ),
        (
            "jockeyProfiles",
            "Trạng thái hồ sơ nghề nghiệp của jockey.",
            [
                ("draft", "Hồ sơ nháp."),
                ("pending", "Chờ duyệt."),
                ("published", "Hồ sơ đã công khai và được dùng trong danh sách jockey."),
                ("rejected", "Hồ sơ bị từ chối."),
                ("archived", "Hồ sơ đã lưu trữ."),
            ],
        ),
        (
            "jockeyTournamentRegistrations",
            "Trạng thái Jockey đăng ký tham gia Tournament.",
            [
                ("pending", "Jockey gửi đăng ký tham gia giải, chờ Admin duyệt."),
                ("approved", "Admin duyệt Jockey tham gia giải."),
                ("rejected", "Admin từ chối Jockey tham gia giải."),
            ],
        ),
        (
            "jockeyInvitations.status",
            "Trạng thái lời mời Owner gửi cho Jockey.",
            [
                ("pending", "Owner gửi lời mời, chờ Jockey phản hồi."),
                ("accepted", "Jockey đồng ý tham gia."),
                ("rejected", "Jockey từ chối."),
                ("cancelled", "Lời mời bị hủy."),
            ],
        ),
        (
            "jockeyInvitations.adminStatus",
            "Trạng thái Admin duyệt cặp Horse + Jockey.",
            [
                ("null", "Chưa tới bước Admin duyệt."),
                ("pending", "Chờ Admin duyệt cặp đăng ký race."),
                ("approved", "Admin duyệt cặp Horse + Jockey."),
                ("rejected", "Admin từ chối cặp Horse + Jockey."),
            ],
        ),
        (
            "races.status",
            "Trạng thái vòng đời của Race.",
            [
                ("draft", "Bản nháp, có trong schema nhưng luồng tạo hiện mở đăng ký ngay."),
                ("registration-open", "Race đang mở đăng ký."),
                ("registration-closed", "Đã đóng đăng ký, chuẩn bị gate/handicap."),
                ("published", "Race đã công bố cho Referee/Live Race."),
                ("in-progress", "Race đang chạy."),
                ("finished", "Kết quả đã được Admin xác nhận."),
                ("completed", "Race hoàn tất award/kết thúc."),
            ],
        ),
        (
            "races.resultStatus",
            "Trạng thái kết quả cấp race.",
            [
                ("draft", "Chưa có kết quả."),
                ("submitted", "Referee đã nộp kết quả, chờ Admin xác nhận."),
                ("approved", "Admin xác nhận kết quả."),
                ("rejected", "Admin từ chối kết quả, cần sửa/nộp lại."),
            ],
        ),
        (
            "raceEntries.status",
            "Trạng thái entry Horse + Jockey trong Race.",
            [
                ("pending-approval", "Entry chờ Admin duyệt."),
                ("approved", "Entry được tham gia race."),
                ("rejected", "Entry bị từ chối."),
            ],
        ),
        (
            "raceEntries.preRaceStatus",
            "Trạng thái kiểm tra trước khi race bắt đầu.",
            [
                ("pending", "Chưa kiểm tra."),
                ("ready-for-referee", "Admin đã chuẩn bị gate/handicap, chờ Referee kiểm tra."),
                ("ready", "Referee xác nhận sẵn sàng."),
                ("absent", "Vắng mặt; khi start race sẽ bị disqualified."),
            ],
        ),
        (
            "raceEntries.resultStatus",
            "Trạng thái kết quả của từng entry.",
            [
                ("draft", "Chưa có kết quả chính thức."),
                ("official", "Kết quả chính thức sau khi Admin xác nhận."),
            ],
        ),
        (
            "raceRefereeAssignments",
            "Trạng thái phân công trọng tài.",
            [
                ("assigned", "Referee đã được phân công."),
                ("confirmed", "Referee xác nhận nhận nhiệm vụ."),
                ("declined", "Referee từ chối nhiệm vụ."),
                ("removed", "Referee bị gỡ khỏi race."),
            ],
        ),
        (
            "refereeReports",
            "Trạng thái báo cáo sự cố của referee.",
            [
                ("draft", "Bản nháp báo cáo."),
                ("submitted", "Đã gửi báo cáo."),
                ("reviewed", "Admin đã xem xét."),
                ("dismissed", "Báo cáo bị bỏ qua/không xử lý tiếp."),
            ],
        ),
        (
            "notifications",
            "Thông báo không có status riêng.",
            [
                ("isRead = false", "Thông báo chưa đọc."),
                ("isRead = true", "Thông báo đã đọc."),
                ("type", "Nhóm thông báo: general, invitation, registration, result, warning."),
            ],
        ),
        (
            "sessions",
            "Session không có status riêng.",
            [
                ("expiresAt", "Còn hiệu lực nếu thời gian hiện tại chưa vượt quá expiresAt."),
            ],
        ),
    ]

    for heading, description, status_rows in sections:
        if heading in {"races.status", "raceEntries.status", "notifications"}:
            doc.add_page_break()

        doc.add_heading(heading, level=1)
        paragraph = doc.add_paragraph(description)
        style_paragraph(paragraph, size=10, color="4B5563")
        add_table(doc, [("Status", "Ý nghĩa")] + status_rows)

    doc.add_section(WD_SECTION.CONTINUOUS)
    note = doc.add_paragraph(
        "Ghi chú: một số trạng thái có trong schema để mở rộng nhưng hiện ít dùng trong luồng demo. "
        "Các trạng thái quan trọng nhất cho bài là users, tournaments, horses, jockeyInvitations, races và raceEntries."
    )
    style_paragraph(note, size=10, color="4B5563")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
