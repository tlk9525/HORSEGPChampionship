from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path("/Users/tranlekhoa/Downloads/Horse_Racing_Project_Thuyet_Trinh.docx")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width):
    cell.width = Inches(width)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def format_table(table, widths=None):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row_index, row in enumerate(table.rows):
        for cell_index, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if widths:
                set_cell_width(cell, widths[cell_index])
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(9)
            if row_index == 0:
                set_cell_shading(cell, "E8EEF5")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def add_heading(doc, text, level=1):
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "Calibri"
        if level == 1:
            run.font.color.rgb = RGBColor(46, 116, 181)
        elif level == 2:
            run.font.color.rgb = RGBColor(31, 77, 120)
    return paragraph


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def add_step_list(doc, items):
    for index, text in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p.add_run(f"{index}. ").bold = True
        p.add_run(text)


def add_code_block(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    for index, line in enumerate(lines):
        if index:
            p.add_run("\n")
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
    return p


def add_kv_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Mục"
    table.rows[0].cells[1].text = "Nội dung"
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = value
    format_table(table, widths=[1.8, 4.5])
    return table


def add_api_table(doc, title, rows):
    add_heading(doc, title, 3)
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "Method"
    hdr[1].text = "Endpoint"
    hdr[2].text = "Mục đích"
    hdr[3].text = "Code liên quan"
    for method, endpoint, purpose, code in rows:
        cells = table.add_row().cells
        cells[0].text = method
        cells[1].text = endpoint
        cells[2].text = purpose
        cells[3].text = code
    format_table(table, widths=[0.8, 2.1, 2.2, 1.2])


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Hệ thống quản lý giải đua ngựa")
    title_run.bold = True
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Tài liệu thuyết trình project, code, API và luồng nghiệp vụ")
    subtitle_run.font.size = Pt(13)
    subtitle_run.font.color.rgb = RGBColor(85, 85, 85)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Project path: ").bold = True
    meta.add_run("/Users/tranlekhoa/Downloads/Horse Racing Tournament Website")

    add_heading(doc, "1. Kết luận ngắn để mở bài", 1)
    doc.add_paragraph(
        "Project này là một ứng dụng web quản lý giải đua ngựa theo vai trò. "
        "Điểm chính của hệ thống là không chỉ hiển thị cuộc đua, mà còn quản lý đầy đủ vòng đời: "
        "tạo giải, tạo race, owner đăng ký ngựa, jockey xác nhận, admin duyệt, referee kiểm tra và ghi kết quả, "
        "sau đó admin duyệt kết quả chính thức để cập nhật rating."
    )
    add_bullet(doc, "Frontend dùng React, TypeScript, Vite và Tailwind CSS để xây giao diện theo vai trò.")
    add_bullet(doc, "Backend dùng Node.js ESM và Hono để chia API theo domain/role.")
    add_bullet(doc, "Database dùng PostgreSQL, schema đặt trong database/postgres/schema.sql.")
    add_bullet(doc, "Xác thực dùng session cookie HttpOnly, an toàn hơn lưu token trong localStorage.")
    add_bullet(doc, "Live race dùng Server-Sent Events để frontend nhận cập nhật theo thời gian thực.")

    add_heading(doc, "2. Tổng quan chức năng", 1)
    add_kv_table(
        doc,
        [
            ("Tên project", "Horse Racing Tournament Website"),
            ("Bài toán", "Quản lý giải đua ngựa, đăng ký ngựa/jockey, kiểm tra trước race, ghi và duyệt kết quả."),
            ("Người dùng chính", "Admin, Owner, Jockey, Referee, Spectator."),
            ("Dữ liệu quan trọng", "users, tournaments, horses, races, raceEntries, registrations, sessions, notifications."),
            ("Điểm nổi bật", "Phân quyền rõ, workflow nhiều bước, rating/handicap, SSE live update, PostgreSQL."),
        ],
    )

    add_heading(doc, "3. Công nghệ sử dụng", 1)
    tech = doc.add_table(rows=1, cols=3)
    tech.rows[0].cells[0].text = "Tầng"
    tech.rows[0].cells[1].text = "Công nghệ"
    tech.rows[0].cells[2].text = "Vai trò"
    for row in [
        ("Frontend", "React 18, TypeScript, Vite", "Xây SPA, route theo trang, gọi API."),
        ("UI", "Tailwind CSS, Lucide React, Recharts", "Style nhanh, icon, biểu đồ/thống kê."),
        ("Backend", "Node.js ESM, Hono", "API server, middleware, route theo vai trò."),
        ("Database", "PostgreSQL", "Lưu user, giải, race, đăng ký, kết quả, session."),
        ("Auth", "bcryptjs, HttpOnly cookie", "Hash password và lưu session an toàn."),
        ("Realtime", "Server-Sent Events", "Đẩy cập nhật live race từ backend sang browser."),
    ]:
        cells = tech.add_row().cells
        cells[0].text, cells[1].text, cells[2].text = row
    format_table(tech, widths=[1.25, 2.3, 2.75])

    add_heading(doc, "4. Kiến trúc hệ thống", 1)
    doc.add_paragraph("Luồng dữ liệu chính khi người dùng thao tác:")
    add_step_list(doc, [
        "Browser tải React frontend.",
        "Frontend gọi /api/bootstrap để lấy dữ liệu khởi động phù hợp với vai trò hiện tại.",
        "Người dùng thao tác trên UI: tạo giải, đăng ký ngựa, duyệt, check-in, ghi kết quả.",
        "Frontend gọi API Hono bằng fetch, kèm credentials để gửi session cookie.",
        "Backend kiểm tra role, validate dữ liệu, ghi PostgreSQL và trả JSON.",
        "Khi race thay đổi, backend phát sự kiện SSE để màn live race cập nhật.",
    ])
    add_code_block(
        doc,
        [
            "Browser / React",
            "   -> frontend/src/app/services/api.ts",
            "   -> /api/*",
            "   -> backend/src/app.js",
            "   -> backend/src/routes/*.js",
            "   -> backend/src/services/*.js",
            "   -> backend/src/sqlDb.js",
            "   -> PostgreSQL",
        ],
    )

    add_heading(doc, "5. Cấu trúc thư mục cần nhớ khi thuyết trình", 1)
    add_kv_table(
        doc,
        [
            ("frontend/src/App.tsx", "Shell chính: auth guard, layout tổng, truyền currentUser xuống route."),
            ("frontend/src/app/AppRoutes.tsx", "Khai báo route React Router cho login, tournaments, races, admin, live race."),
            ("frontend/src/app/routing.ts", "Helper điều hướng: map role về home page và map path về page hiện tại."),
            ("frontend/src/app/services/api.ts", "Client API contract: type dữ liệu, request helper, login/register/bootstrap/admin/owner/referee calls."),
            ("backend/src/app.js", "Khởi tạo Hono app, middleware security/CORS/body limit, mount các route /api."),
            ("backend/src/routes/*.js", "API chia theo public/auth/admin/owner/jockey/referee/notifications."),
            ("backend/src/services/handicapService.js", "Logic rating, handicap, eligibility range và post-race rating."),
            ("database/postgres/schema.sql", "Nguồn chính cho database schema PostgreSQL."),
        ],
    )

    add_heading(doc, "6. Phân quyền và vai trò", 1)
    role_table = doc.add_table(rows=1, cols=2)
    role_table.rows[0].cells[0].text = "Vai trò"
    role_table.rows[0].cells[1].text = "Quyền chính"
    for role, rights in [
        ("Admin", "Tạo giải, tạo race, phân công referee, duyệt tài khoản/hồ sơ/đăng ký, bắt đầu/kết thúc race, duyệt kết quả."),
        ("Owner", "Tạo hồ sơ ngựa, đăng ký ngựa vào race, chọn hoặc mời jockey."),
        ("Jockey", "Tạo hồ sơ jockey, đăng ký khả dụng, chấp nhận/từ chối lời mời."),
        ("Referee", "Xem race được phân công, check-in entry, ghi vị trí/thời gian/vi phạm, nộp kết quả nháp."),
        ("Spectator", "Xem giải, race card, live race và kết quả công khai."),
    ]:
        cells = role_table.add_row().cells
        cells[0].text = role
        cells[1].text = rights
    format_table(role_table, widths=[1.2, 5.1])

    add_heading(doc, "7. Luồng nghiệp vụ chính", 1)
    add_step_list(doc, [
        "Admin tạo tournament.",
        "Admin tạo các race thuộc tournament, đặt ngày giờ, địa điểm, hạng đua, rating range, handicap range và referee.",
        "Owner tạo hồ sơ ngựa, hồ sơ cần được duyệt trước khi tham gia race.",
        "Jockey tạo hồ sơ và đăng ký khả dụng hoặc nhận lời mời từ owner.",
        "Owner đăng ký ngựa vào race đang mở.",
        "Admin duyệt đăng ký/cặp horse-jockey.",
        "Khi đóng đăng ký, hệ thống snapshot rating, tính handicap và xếp lane.",
        "Admin publish race để công bố danh sách xuất phát.",
        "Referee kiểm tra trước race: ready, absent, incident hoặc scratched.",
        "Admin start race, sau đó finish race.",
        "Referee nhập kết quả nháp và submit.",
        "Admin approve kết quả chính thức, hệ thống cập nhật rating và chuyển race sang completed.",
    ])

    add_heading(doc, "8. Các trạng thái quan trọng", 1)
    status_table = doc.add_table(rows=1, cols=2)
    status_table.rows[0].cells[0].text = "Nhóm trạng thái"
    status_table.rows[0].cells[1].text = "Giá trị trong code/database"
    for group, values in [
        ("Race lifecycle", "draft, registration-open, registration-closed, published, in-progress, finished, completed, cancelled"),
        ("Đăng ký ngựa vào race", "pending-jockey, pending-admin, approved, rejected, cancelled"),
        ("Entry xuất phát", "approved, scratched"),
        ("Pre-race check", "pending, ready, absent, incident, scratched"),
        ("Kết quả", "draft, submitted, official, disqualified"),
    ]:
        cells = status_table.add_row().cells
        cells[0].text = group
        cells[1].text = values
    format_table(status_table, widths=[2.0, 4.3])

    add_heading(doc, "9. Giải thích frontend", 1)
    doc.add_paragraph(
        "Frontend là Single Page Application. Người dùng chuyển trang bằng React Router, còn dữ liệu được lấy qua các hàm trong services/api.ts."
    )
    add_bullet(doc, "App.tsx giữ trạng thái user hiện tại, xử lý login/logout và bảo vệ route.")
    add_bullet(doc, "AppRoutes.tsx khai báo các route như /tournaments, /races/:raceId, /races/:raceId/register, /admin/races/new, /live-race/:raceId.")
    add_bullet(doc, "services/api.ts định nghĩa type TypeScript cho User, Horse, Race, Tournament, RaceEntry và các hàm gọi API.")
    add_bullet(doc, "getBootstrap() có cache ngắn 10 giây để đổi trang không gọi endpoint nặng liên tục.")
    add_bullet(doc, "Các request POST/PATCH/DELETE sẽ invalidate cache để dữ liệu mới được tải lại.")
    add_code_block(
        doc,
        [
            "const response = await fetch(`${API_URL}${path}`, {",
            "  credentials: 'include',",
            "  headers: { 'Content-Type': 'application/json' },",
            "});",
        ],
    )

    add_heading(doc, "10. Giải thích backend", 1)
    doc.add_paragraph("Backend dùng Hono, chia route theo vai trò để dễ bảo trì và dễ kiểm soát quyền truy cập.")
    add_bullet(doc, "app.js mount /api public, auth, owner, jockey, admin, referee và notifications.")
    add_bullet(doc, "secureHeaders() tăng bảo mật HTTP header.")
    add_bullet(doc, "bodyLimit giới hạn request body tối đa 1MB.")
    add_bullet(doc, "CORS cho phép frontend local và FRONTEND_URL gửi request kèm cookie.")
    add_bullet(doc, "Route admin dùng middleware requireRole(..., ['admin']) nên user không phải admin nhận 403.")
    add_code_block(
        doc,
        [
            "app.route('/api', createPublicRoutes(getDb));",
            "app.route('/api', createAuthRoutes(...));",
            "app.route('/api/owner', createOwnerRoutes(getDb, writeDb));",
            "app.route('/api/admin', createAdminRoutes(getDb, writeDb, persistAdminRaceAction));",
            "app.route('/api/referee', createRefereeRoutes(...));",
        ],
    )

    add_heading(doc, "11. Auth và bảo mật", 1)
    add_bullet(doc, "Password được hash bằng bcryptjs khi đăng ký hoặc khi login gặp password cũ chưa hash.")
    add_bullet(doc, "Session token tạo bằng randomUUID và lưu trong bảng sessions.")
    add_bullet(doc, "Cookie dùng httpOnly nên JavaScript phía browser không đọc trực tiếp được token.")
    add_bullet(doc, "Các role cần phê duyệt sẽ có status pending và không đăng nhập được cho tới khi admin duyệt.")
    add_bullet(doc, "Backend trả publicUser để không gửi password ra frontend.")

    add_heading(doc, "12. Rating và handicap", 1)
    doc.add_paragraph("Logic chính nằm trong backend/src/services/handicapService.js.")
    add_bullet(doc, "Rating ban đầu = speed 35% + stamina 25% + form 30% + health 10%.")
    add_bullet(doc, "officialHorseRating ưu tiên overallRating đã lưu; nếu chưa có thì tính từ các chỉ số hồ sơ.")
    add_bullet(doc, "raceEligibilityRange kiểm tra ratingMin/ratingMax của race hoặc fallback theo class.")
    add_bullet(doc, "computeRaceHandicap tính trọng lượng chỉ định theo rating cao nhất trong field.")
    add_bullet(doc, "computePostRaceRating so sánh kết quả thực tế với điểm kỳ vọng, clamp thay đổi trong khoảng -8 đến +8.")
    add_code_block(
        doc,
        [
            "Assigned weight = Handicap Max - (Highest Field Rating - Horse Rating)",
            "Post Race Rating = clamp(Rating Snapshot + Rating Change, 0, 140)",
            "Rating Change = round(clamp(10 * (Actual - Expected) * FieldFactor, -8, +8))",
        ],
    )

    add_heading(doc, "13. Database schema", 1)
    db_table = doc.add_table(rows=1, cols=3)
    db_table.rows[0].cells[0].text = "Bảng"
    db_table.rows[0].cells[1].text = "Ý nghĩa"
    db_table.rows[0].cells[2].text = "Quan hệ chính"
    for row in [
        ("users", "Tài khoản và role", "Owner/Jockey/Referee/Admin liên kết tới nhiều bảng khác."),
        ("tournaments", "Giải đấu", "Một tournament có nhiều races."),
        ("horses", "Hồ sơ ngựa", "Mỗi horse thuộc một ownerUserId."),
        ("races", "Cuộc đua", "Thuộc tournament, có trạng thái, thời gian, class, handicap."),
        ("raceRefereeAssignments", "Phân công trọng tài", "Nối race với user role referee."),
        ("jockeyProfiles", "Hồ sơ jockey", "Mỗi jockey user có một profile."),
        ("jockeyRaceRegistrations", "Jockey đăng ký khả dụng", "Theo từng race."),
        ("jockeyInvitations", "Owner mời jockey", "Nối horse, owner, jockey, race."),
        ("horseRaceRegistrations", "Đăng ký ngựa vào race", "Bước chờ jockey/admin duyệt trước khi thành entry."),
        ("raceEntries", "Danh sách xuất phát chính thức", "Chứa lane, handicap, rating snapshot, result."),
        ("notifications", "Thông báo", "Gửi cho user khi có duyệt hoặc thay đổi trạng thái."),
        ("sessions", "Phiên đăng nhập", "Cookie token map tới user."),
    ]:
        cells = db_table.add_row().cells
        cells[0].text, cells[1].text, cells[2].text = row
    format_table(db_table, widths=[1.7, 2.25, 2.35])

    add_heading(doc, "14. API chính", 1)
    add_api_table(
        doc,
        "Auth",
        [
            ("GET", "/api/me", "Lấy user hiện tại từ session cookie.", "authRoutes.js"),
            ("POST", "/api/login", "Đăng nhập, tạo session, set HttpOnly cookie.", "authRoutes.js"),
            ("POST", "/api/register", "Đăng ký tài khoản theo role được phép.", "authRoutes.js"),
            ("POST", "/api/logout", "Xóa session và cookie.", "authRoutes.js"),
        ],
    )
    add_api_table(
        doc,
        "Public và dữ liệu chung",
        [
            ("GET", "/api/health", "Health check backend.", "publicRoutes.js"),
            ("GET", "/api/bootstrap", "Trả dữ liệu tổng hợp theo role.", "publicRoutes.js"),
            ("GET", "/api/live/races/:id/events", "SSE stream cho live race.", "publicRoutes.js"),
        ],
    )
    add_api_table(
        doc,
        "Admin",
        [
            ("GET", "/api/admin/approvals", "Danh sách yêu cầu chờ duyệt.", "adminRoutes.js"),
            ("POST", "/api/admin/approvals/:type/:id", "Duyệt hoặc từ chối account/horse/race registration/pairing.", "adminRoutes.js"),
            ("POST", "/api/admin/tournaments", "Tạo giải đấu.", "adminRoutes.js"),
            ("PATCH", "/api/admin/tournaments/:id", "Cập nhật giải đấu.", "adminRoutes.js"),
            ("DELETE", "/api/admin/tournaments/:id", "Xóa giải và dữ liệu liên quan.", "adminRoutes.js"),
            ("GET", "/api/admin/race-builder", "Lấy dữ liệu tạo race: tournaments, races, referees.", "adminRoutes.js"),
            ("POST", "/api/admin/races", "Tạo race mới.", "adminRoutes.js"),
            ("PATCH", "/api/admin/races/:id", "Sửa race chưa công bố.", "adminRoutes.js"),
            ("DELETE", "/api/admin/races/:id", "Xóa race chưa công bố.", "adminRoutes.js"),
            ("POST", "/api/admin/races/:id/:action", "close-registration, publish, start-race, finish-race, complete-results, cancel-race.", "adminRoutes.js"),
        ],
    )
    add_api_table(
        doc,
        "Owner, Jockey, Referee",
        [
            ("GET", "/api/owner/portal", "Dữ liệu dashboard owner.", "ownerRoutes.js"),
            ("GET", "/api/owner/race-registration?raceId=...", "Dữ liệu form đăng ký race.", "ownerRoutes.js"),
            ("POST", "/api/owner/horses", "Tạo hồ sơ ngựa.", "ownerRoutes.js"),
            ("POST", "/api/owner/race-registrations", "Đăng ký ngựa hoặc gửi lời mời jockey.", "ownerRoutes.js"),
            ("GET", "/api/jockey/portal", "Dữ liệu dashboard jockey.", "jockeyRoutes.js"),
            ("POST", "/api/jockey/profile", "Lưu hồ sơ jockey.", "jockeyRoutes.js"),
            ("POST", "/api/jockey/race-registrations", "Jockey đăng ký tham gia race.", "jockeyRoutes.js"),
            ("POST", "/api/jockey/invitations/:id", "Jockey nhận/từ chối lời mời.", "jockeyRoutes.js"),
            ("POST", "/api/referee/races/:id/submit-results", "Nộp kết quả nháp.", "refereeRoutes.js"),
            ("POST", "/api/referee/race-entries/:id/readiness/:status", "Đánh dấu ready/absent/incident/scratched.", "refereeRoutes.js"),
            ("POST", "/api/referee/race-entries/:id/result", "Ghi vị trí và thời gian hoàn thành.", "refereeRoutes.js"),
        ],
    )

    add_heading(doc, "15. Cách chạy project để demo", 1)
    add_code_block(
        doc,
        [
            "npm install",
            "npm run db:init",
            "npm run api",
            "npm run dev",
        ],
    )
    doc.add_paragraph("URL mặc định:")
    add_bullet(doc, "Backend: http://127.0.0.1:4000")
    add_bullet(doc, "Frontend: http://127.0.0.1:5173")
    add_bullet(doc, "Health check: http://127.0.0.1:4000/api/health")
    doc.add_paragraph("Kiểm tra toàn bộ trước khi nộp/thuyết trình:")
    add_code_block(doc, ["npm run check"])

    add_heading(doc, "16. Kịch bản demo đề xuất", 1)
    add_step_list(doc, [
        "Mở trang login và đăng nhập admin.",
        "Vào admin panel, tạo tournament hoặc mở tournament có sẵn.",
        "Tạo race mới với rating range, handicap range và referee.",
        "Đăng nhập owner, tạo hồ sơ ngựa và đăng ký vào race.",
        "Đăng nhập jockey, tạo profile và chấp nhận lời mời hoặc đăng ký race.",
        "Quay lại admin duyệt các yêu cầu.",
        "Đóng đăng ký, publish race và giải thích hệ thống snapshot rating/handicap/lane.",
        "Đăng nhập referee để đánh dấu ready và nhập kết quả.",
        "Admin duyệt kết quả chính thức, xem rating thay đổi và race completed.",
        "Mở live race để cho thấy SSE/live replay.",
    ])

    add_heading(doc, "17. Những điểm nên nhấn mạnh khi bị hỏi", 1)
    qa = doc.add_table(rows=1, cols=2)
    qa.rows[0].cells[0].text = "Câu hỏi"
    qa.rows[0].cells[1].text = "Câu trả lời ngắn"
    for question, answer in [
        ("Vì sao dùng Hono?", "Hono nhẹ, route rõ, middleware đơn giản, phù hợp API Node.js nhỏ và vừa."),
        ("Vì sao dùng HttpOnly cookie?", "Giảm rủi ro token bị đọc bởi JavaScript nếu có XSS."),
        ("Vì sao có /api/bootstrap?", "Frontend cần nhiều dữ liệu ban đầu; gom lại giúp giảm số request và lọc theo role ở backend."),
        ("Vì sao tách raceRegistrations và raceEntries?", "registration là yêu cầu chờ duyệt; raceEntries là danh sách xuất phát chính thức sau duyệt."),
        ("Rating có phải chuẩn HKJC không?", "Không. Đây là công thức nội bộ để demo logic rating/handicap trong project."),
        ("SSE dùng để làm gì?", "Giúp màn live race nhận cập nhật realtime mà không cần polling liên tục."),
        ("Điểm mạnh của project?", "Có workflow thực tế, phân quyền, database rõ, bảo mật session, nghiệp vụ rating/handicap và live update."),
        ("Rủi ro/cải tiến?", "Có thể thêm test E2E, audit log rộng hơn, OpenAPI docs, rate limit, refresh token/session cleanup job."),
    ]:
        cells = qa.add_row().cells
        cells[0].text = question
        cells[1].text = answer
    format_table(qa, widths=[2.0, 4.3])

    add_heading(doc, "18. Phần kết luận", 1)
    doc.add_paragraph(
        "Tóm lại, project thể hiện một hệ thống web full-stack có luồng nghiệp vụ rõ, không chỉ CRUD đơn giản. "
        "Phần đáng chú ý nhất là cách hệ thống chia quyền theo vai trò, tách giai đoạn đăng ký với danh sách xuất phát, "
        "tính handicap/rating bằng service riêng và dùng SSE cho live race. Khi thuyết trình, nên đi theo luồng nghiệp vụ trước, "
        "sau đó mới chỉ vào code frontend, backend, database và API."
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "Phụ lục A. File code nguồn nên mở khi demo", 1)
    for item in [
        "frontend/src/app/services/api.ts: endpoint contract giữa frontend và backend.",
        "backend/src/app.js: nơi mount middleware và route.",
        "backend/src/routes/adminRoutes.js: workflow admin và race lifecycle.",
        "backend/src/routes/authRoutes.js: login/register/logout/session cookie.",
        "backend/src/routes/publicRoutes.js: bootstrap, health và SSE.",
        "backend/src/services/handicapService.js: rating và handicap.",
        "database/postgres/schema.sql: database schema.",
        "docs/business-flow-and-roles.md: mô tả workflow và role.",
        "docs/schema-erd.md: ERD dạng Mermaid.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Phụ lục B. Lệnh nhanh", 1)
    add_code_block(
        doc,
        [
            "# Cài thư viện",
            "npm install",
            "",
            "# Tạo database schema + seed",
            "npm run db:init",
            "",
            "# Chạy backend",
            "npm run api",
            "",
            "# Chạy frontend",
            "npm run dev",
            "",
            "# Kiểm tra project",
            "npm run check",
        ],
    )

    doc.core_properties.title = "Horse Racing Project - Tài liệu thuyết trình"
    doc.core_properties.subject = "Project, code, API, database, workflow"
    doc.core_properties.author = "Codex"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
